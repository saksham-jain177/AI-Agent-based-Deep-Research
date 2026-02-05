import streamlit as st
from main import run_research  # Import run_research from main.py
from draft_agent import format_citation, STYLE_TEMPLATES  # Add this import
from cost_estimator import estimate_research_cost  # Cost estimation
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT
import io
import json
import html
import requests
import datetime
import logging
import re
import os
import time
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches

# Set up logging
logging.basicConfig(filename="research_agent.log", level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

# Initialize all session state variables at the top
if "research_data" not in st.session_state:
    st.session_state.research_data = None
if "response" not in st.session_state:
    st.session_state.response = None
if "pdf_buffer" not in st.session_state:
    st.session_state.pdf_buffer = None
if "word_buffer" not in st.session_state:
    st.session_state.word_buffer = None
if "writing_style" not in st.session_state:
    st.session_state.writing_style = "Academic"
if "language" not in st.session_state:
    st.session_state.language = "English"
if "citation_format" not in st.session_state:
    st.session_state.citation_format = "APA"
if "target_word_count" not in st.session_state:
    st.session_state.target_word_count = 1000  # Default value
if "__busy__" not in st.session_state:
    st.session_state["__busy__"] = False
if "__trigger__" not in st.session_state:
    st.session_state["__trigger__"] = False

# Inject custom CSS for improved readability and aesthetics
st.markdown("""
    <style>
    /* Base styling - Futuristic Dark Theme */
    .stApp {
        background-color: #0d1117;
        color: #e6edf3;
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
        line-height: 1.6;
    }
    
    /* Typography */
    h1 {
        color: #58a6ff;
        font-size: 2.5rem;
        font-weight: 700;
        letter-spacing: -0.02em;
        margin-bottom: 1.5rem;
        padding-bottom: 0.6rem; /* Removed bottom border to avoid double line */
        text-shadow: 0 0 15px rgba(88, 166, 255, 0.5);
    }
    
    h2 {
        color: #79c0ff;
        font-size: 1.8rem;
        font-weight: 600;
        margin-top: 1.8rem;
        margin-bottom: 1rem;
    }
    
    h3 {
        color: #a5d6ff;
        font-size: 1.4rem;
        font-weight: 500;
        margin-top: 1.5rem;
        margin-bottom: 0.8rem;
    }
    
    p, li {
        font-size: 1rem;
        color: #c9d1d9;
        line-height: 1.7;
    }
    
    /* Text input styling */
    .stTextInput > div > input {
        background-color: #161b22;
        color: #c9d1d9;
        border: 1px solid #30363d;
        border-radius: 8px;
        box-shadow: 0 0 5px rgba(31, 111, 235, 0.2);
        padding: 12px 16px;
        font-size: 1rem;
        transition: all 0.2s ease;
    }
    
    .stTextInput > div > input:focus {
        border-color: #1f6feb;
        box-shadow: 0 0 10px rgba(31, 111, 235, 0.4);
        outline: none;
    }
    
    /* Button styling (Clean & Modern) */
    .stButton > button {
        background-color: #1f6feb;
        color: white;
        font-weight: 600;
        border-radius: 6px;
        padding: 0.5rem 1.5rem;
        border: 1px solid rgba(255,255,255,0.1);
        transition: background-color 0.2s;
        overflow: hidden; /* Prevent pseudo-elements from leaking into background */
        position: relative;
    }
    
    .stButton > button:hover {
        background-color: #388bfd;
    }
    
    .stButton > button:disabled {
        background-color: #21262d !important;
        color: #8b949e !important;
        border-color: #30363d !important;
    }
    
    .stButton > button::after {
        content: '';
        position: absolute;
        top: -50%;
        left: -50%;
        width: 200%;
        height: 200%;
        background: rgba(255, 255, 255, 0.1);
        transform: rotate(30deg);
        transition: transform 0.3s ease;
    }
    
    .stButton > button:hover::after {
        transform: rotate(30deg) translate(-10%, -10%);
    }
    
    /* Download button styling (for Downloading PDF... üìÑ button) */
    .stDownloadButton > button {
        background: linear-gradient(45deg, #065f46, #0ea5e9);
        color: white;
        font-weight: 600;
        padding: 0.7rem 2rem;
        border-radius: 8px;
        border: none;
        cursor: pointer;
        transition: all 0.3s ease;
        box-shadow: 0 0 15px rgba(14, 165, 233, 0.5);
        position: relative;
        overflow: hidden; /* Prevent pseudo-elements from leaking into background */
        display: inline-flex;
        align-items: center;
        gap: 0.5rem;
        font-size: 1rem;
        margin-top: 0.5rem;
    }
    
    .stDownloadButton > button:hover {
        box-shadow: 0 0 25px rgba(14, 165, 233, 0.7);
        background: linear-gradient(45deg, #0ea5e9, #065f46);
    }
    
    .stDownloadButton > button:active {
        box-shadow: 0 0 10px rgba(14, 165, 233, 0.3);
    }
    
    .stDownloadButton > button::after {
        content: '';
        position: absolute;
        top: -50%;
        left: -50%;
        width: 200%;
        height: 200%;
        background: rgba(255, 255, 255, 0.1);
        transform: rotate(30deg);
        transition: transform 0.3s ease;
    }
    
    .stDownloadButton > button:hover::after {
        transform: rotate(30deg) translate(-10%, -10%);
    }
    
    /* Custom Selectbox styling */
    .custom-select-wrapper {
        position: relative;
        width: 100%;
        font-size: 1rem;
    }
    
    .custom-select {
        background: linear-gradient(45deg, #065f46, #0ea5e9);
        color: white;
        font-weight: 600;
        border: none;
        border-radius: 8px;
        padding: 10px;
        font-size: 1rem;
        transition: all 0.3s ease;
        box-shadow: 0 0 10px rgba(14, 165, 233, 0.3);
        width: 100%;
        cursor: pointer;
        appearance: none;
        -webkit-appearance: none;
        -moz-appearance: none;
    }
    
    .custom-select:hover {
        box-shadow: 0 0 15px rgba(14, 165, 233, 0.5);
        transform: translateY(-2px) scale(1.02);
        background: linear-gradient(45deg, #0ea5e9, #065f46);
    }
    
    .custom-select:focus {
        outline: none;
        box-shadow: 0 0 15px rgba(14, 165, 233, 0.5);
    }
    
    /* Custom arrow for the dropdown */
    .custom-select-wrapper::after {
        content: '\\25BC'; /* Unicode for down arrow */
        position: absolute;
        right: 15px;
        top: 50%;
        transform: translateY(-50%);
        color: white;
        font-size: 1rem;
        pointer-events: none;
    }
    
    /* Dropdown menu styling */
    .custom-select option {
        background: #161b22;
        color: #ffffff;
        font-weight: 500;
        border: 1px solid #1f6feb;
        border-radius: 8px;
        box-shadow: 0 0 5px rgba(14, 165, 233, 0.2);
        padding: 10px;
    }
    
    .custom-select option:hover {
        background: linear-gradient(45deg, #0ea5e9, #065f46);
        color: white;
        box-shadow: 0 0 10px rgba(14, 165, 233, 0.3);
    }
    
    /* Ensure selectboxes are visible (we use native Streamlit controls) */
    .stSelectbox { display: block !important; }
    
    /* Sidebar styling */
    .stSidebar {
        background-color: #161b22;
        border-right: 1px solid #30363d;
    }
    
    .stSidebar .stMarkdown {
        color: #c9d1d9 !important;
    }
    
    .stSidebar h2 {
        color: #58a6ff;
        font-size: 1.4rem;
        margin-top: 2rem;
        margin-bottom: 1rem;
        padding-bottom: 0.5rem;
        border-bottom: 1px solid #30363d;
        text-shadow: 0 0 10px rgba(88, 166, 255, 0.3);
    }
    
    /* Status indicators in sidebar */
    .stSidebar .stSuccess {
        background: linear-gradient(90deg, #033a2c, #065f46) !important;
        color: #ecfdf5 !important;
        padding: 10px 15px;
        border-radius: 8px;
        font-weight: 500;
        margin-bottom: 1rem;
        display: flex;
        align-items: center;
        box-shadow: 0 0 10px rgba(6, 95, 70, 0.4);
        border-left: 3px solid #10b981;
    }
    
    .stSidebar .stError {
        background: linear-gradient(90deg, #770b0b, #9a1515) !important;
        color: #fee2e2 !important;
        padding: 10px 15px;
        border-radius: 8px;
        font-weight: 500;
        margin-bottom: 1rem;
        display: flex;
        align-items: center;
        box-shadow: 0 0 10px rgba(185, 28, 28, 0.4);
        border-left: 3px solid #ef4444;
    }
    
    /* JSON/Code block styling */
    .stCodeBlock {
        background-color: #0d1117;
        color: #c9d1d9;
        border-radius: 8px;
        padding: 1.2rem;
        font-family: 'Fira Code', 'JetBrains Mono', 'Consolas', monospace;
        font-size: 0.9rem;
        line-height: 1.6;
        overflow-x: auto;
        border: 1px solid #30363d;
        box-shadow: inset 0 0 10px rgba(0, 0, 0, 0.2);
    }
    
    /* Progress bar - keep Streamlit defaults for correct fill behavior */
    .stProgress > div > div { height: 0.6rem !important; border-radius: 1rem; }
    .stProgress > div > div > div { border-radius: 1rem; }
    
    /* Notification messages */
    .stSuccess {
        background-color: rgba(6, 95, 70, 0.2) !important;
        color: #34d399 !important;
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid #10b981;
        margin: 1rem 0;
        font-size: 1rem;
        box-shadow: 0 0 10px rgba(6, 95, 70, 0.2);
    }
    
    .stInfo {
        background-color: rgba(37, 99, 235, 0.2) !important;
        color: #93c5fd !important;
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid #3b82f6;
        margin: 1rem 0;
        font-size: 1rem;
        box-shadow: 0 0 10px rgba(37, 99, 235, 0.2);
    }
    
    .stWarning {
        background-color: rgba(217, 119, 6, 0.2) !important;
        color: #fbbf24 !important;
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid #f59e0b;
        margin: 1rem 0;
        font-size: 1rem;
        box-shadow: 0 0 10px rgba(217, 119, 6, 0.2);
    }
    
    .stError {
        background-color: rgba(185, 28, 28, 0.2) !important;
        color: #f87171 !important;
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid #ef4444;
        margin: 1rem 0;
        font-size: 1rem;
        box-shadow: 0 0 10px rgba(185, 28, 28, 0.2);
    }
    
    /* JSON display */
    .stJson {
        background-color: #161b22;
        border: 1px solid #30363d;
        border-radius: 8px;
        padding: 1.2rem;
        margin: 1rem 0;
        box-shadow: 0 0 15px rgba(0, 0, 0, 0.2);
    }
    
    /* Download button container */
    .download-btn-container {
        margin-top: 1rem;
        text-align: center;
    }
    
    /* Format description */
    .format-description {
        margin-top: 0.5rem;
        font-size: 0.9rem;
        color: #a5d6ff;
        text-align: center;
        transition: opacity 0.3s ease;
    }
    
    /* Spinner */
    .stSpinner > div {
        border-color: #0ea5e9 !important;
        border-bottom-color: transparent !important;
        filter: drop-shadow(0 0 8px rgba(14, 165, 233, 0.6));
    }
    
    /* Feedback textarea */
    .stTextArea > div > textarea {
        background-color: #161b22;
        color: #c9d1d9;
        border: 1px solid #30363d;
        border-radius: 8px;
        padding: 12px;
        font-size: 0.9rem;
        line-height: 1.5;
        resize: vertical;
        box-shadow: 0 0 5px rgba(31, 111, 235, 0.2);
    }
    
    .stTextArea > div > textarea:focus {
        border-color: #1f6feb;
        box-shadow: 0 0 10px rgba(31, 111, 235, 0.4);
        outline: none;
    }
    
    /* General spacing and containers */
    .main .block-container {
        padding: 2rem 1.5rem;
        max-width: 960px;
    }
    
    .stMarkdown {
        margin-bottom: 1.5rem;
    }
    
    /* Custom glowing effects and futuristic touches */
    .stApp::before {
        content: "";
        position: fixed;
        top: 0;
        left: 0;
        width: 100%;
        height: 100%;
        background: radial-gradient(circle at top right, rgba(14, 165, 233, 0.05), transparent 60%);
        pointer-events: none;
        z-index: -1;
    }
    
    /* Title underline animation */
    h1::after {
        content: '';
        display: block;
        position: relative; /* Changed from absolute to prevent overlap */
        height: 2px;
        width: 100px;
        max-width: 100%;
        background: #1f6feb;
        margin-top: 5px;
        background: linear-gradient(90deg, #1f6feb, transparent);
    }
    
    /* Limit size of headings inside results to prevent "Huge Bold" looks */
    .stExpander h1, .stExpander h2, .stExpander h3 {
        font-size: 1.2rem !important;
        margin-top: 1rem !important;
        margin-bottom: 0.5rem !important;
        border: none !important;
    }
    
    .stExpander strong {
        color: #58a6ff;
    }
    </style>
    """, unsafe_allow_html=True)

def _clean_analysis_text(text):
    """Deep cleaning for Analysis sections to handle LLM artifacts."""
    if not isinstance(text, str):
        return ""
    t = text.strip()
    
    # Aggressive: Strip outer bolding that wraps EVERYTHING
    # Sometimes LLMs do: ** \n content \n **
    if (t.startswith("**") and t.endswith("**")) or (t.startswith("__") and t.endswith("__")):
        # Only strip if it's a long block, otherwise it might be a legitimate single-word bold
        if len(t) > 100:
            t = t[2:-2].strip()

    # Remove redundant headings that repeat the section title (H1, H2, H3)
    # Often the LLM starts with "# Analysis" or "## Discussion"
    t = re.sub(r'^(#|##|###)\s+[^\n]+\n*', '', t) 
    
    # Fix mashed-together tables: "Text | Col1 | Col2 |" -> "Text\n\n| Col1 | Col2 |"
    # Look for sentences ending in punctuation followed by a pipe
    t = re.sub(r'([\.!?„ÄÇÔºÅÔºü])\s*(\|)', r'\1\n\n\2', t)
    
    lines = t.split('\n')
    new_lines = []
    for line in lines:
        line_stripped = line.strip()
        # If a line looks like a table row but has leading text
        if "|" in line_stripped and not line_stripped.startswith("|") and line_stripped.count('|') >= 2:
            parts = re.split(r"(?=\|)", line, 1) # Split at first |
            new_lines.append(parts[0].strip())
            new_lines.append("") # spacer
            new_lines.append(parts[1].strip())
        else:
            new_lines.append(line)
    
    # Repair missing table dividers
    temp_lines = new_lines
    new_lines = []
    i = 0
    while i < len(temp_lines):
        line = temp_lines[i]
        line_stripped = line.strip()
        if line_stripped.startswith('|') and line_stripped.count('|') >= 3:
            # Possible table line
            is_divider = '---' in line_stripped or (line_stripped.count('-') > 5 and '|' in line_stripped)
            new_lines.append(line)
            
            # If this is a header (not a divider) and there's no divider next, inject one
            if not is_divider:
                has_divider = False
                if i + 1 < len(temp_lines):
                    next_line = temp_lines[i+1].strip()
                    if next_line.startswith('|') and ('---' in next_line or '-|' in next_line):
                        has_divider = True
                
                if not has_divider:
                    # Only inject if it looks like a header (3+ pipes)
                    num_col = line_stripped.count('|') - 1
                    divider = "|" + "|".join(["---"] * num_col) + "|"
                    new_lines.append(divider)
        else:
            new_lines.append(line)
        i += 1
        
    t = "\n".join(new_lines)
    
    # Fix BibTeX blocks: isolate them and ensure they are code blocks or clean text
    if "bibtex" in t.lower():
        # If it looks like raw bibtex without code blocks, wrap it
        if "@" in t and "{" in t and "```" not in t:
            t = re.sub(r'(@[a-zA-Z]+\{[^}]+})', r'\n```bibtex\n\1\n```\n', t)

    # Ensure consistent spacing
    t = re.sub(r'\n{3,}', '\n\n', t)
    
    return t.strip()

# Function to check OpenRouter status
def check_openrouter_status():
    """Check if OpenRouter API is operational."""
    try:
        response = requests.get("https://openrouter.ai/api/v1/models", timeout=5)
        return response.status_code == 200
    except Exception as e:
        logging.error(f"Failed to check OpenRouter status: {str(e)}")
        return False

# Function to add page numbers to the PDF
def on_page(canvas, doc):
    page_num = canvas.getPageNumber()
    text = f"Page {page_num}"
    canvas.saveState()
    canvas.setFont('Helvetica', 8)
    canvas.setFillColor(colors.grey)
    canvas.drawRightString(doc.rightMargin + doc.width, doc.bottomMargin - 10, text)
    canvas.restoreState()

def format_reference_for_pdf(ref_text):
    """Format a single reference for PDF."""
    # Extract date and URL using regex
    date_match = re.search(r'\(([0-9]{4},\s*[^)]+)\)', ref_text)
    url_match = re.search(r'Retrieved from\s+(https?://\S+)', ref_text)
    
    if date_match and url_match:
        date = date_match.group(1)
        url = url_match.group(1)
        return f"({date}). Retrieved from {url}"
    return ref_text

def preprocess_references(refs_section):
    """Preprocess references to ensure proper formatting."""
    # Split references by looking for date pattern and "Retrieved from"
    pattern = r'(?=\([0-9]{4},.*?\).*?Retrieved from)'
    refs = re.split(pattern, refs_section)
    # Clean up each reference
    refs = [ref.strip() for ref in refs if ref.strip()]
    
    formatted_refs = []
    for i, ref in enumerate(refs, 1):
        # Split URL if it's broken across lines
        ref = re.sub(r'(?<=https?://\S+)-\s*\n\s*(?=\S+)', '', ref)
        # Remove any "Page X" artifacts
        ref = re.sub(r'Page \d+\s*', '', ref)
        # Clean up extra whitespace
        ref = ' '.join(ref.split())
        formatted_ref = format_reference_for_pdf(ref, i)
        formatted_refs.append(formatted_ref)
    
    return formatted_refs

def format_references_section(refs_text):
    """Format references by adding line breaks after URLs."""
    # Split by "Retrieved from" to separate different references
    refs = refs_text.split("Retrieved from")
    formatted_refs = []
    
    for i, ref in enumerate(refs):
        if i == 0:  # First part might be empty
            continue
            
        # Find the URL pattern
        url_match = re.search(r'(https?://\S+)', ref)
        if url_match:
            url = url_match.group(1)
            # Get the text before the URL (title part)
            title_part = ref[:url_match.start()].strip()
            if i > 0 and formatted_refs:  # Add the previous title
                formatted_refs[-1] += f"Retrieved from {url}\n\n"
            formatted_refs.append(title_part)
    
    # Handle the last URL
    if formatted_refs and url_match:
        formatted_refs[-1] += f"Retrieved from {url}\n\n"
    
    return "".join(formatted_refs)

# Function to generate PDF with proper formatting and cover page
def generate_pdf(query, data, summary, deep_research=False):
    """Generate a PDF report with query, data, and summary in a research paper format."""
    # Parse JSON if needed
    parsed = None
    if isinstance(summary, str) and summary.strip().startswith('{'):
        try:
            parsed = json.loads(summary)
        except:
            pass
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        topMargin=72,
        bottomMargin=72,
        leftMargin=72,
        rightMargin=72
    )
    styles = getSampleStyleSheet()

    # Customize styles for a research paper look
    styles['Title'].fontSize = 16
    styles['Title'].spaceAfter = 12
    styles['Heading2'].fontSize = 14
    styles['Heading2'].spaceAfter = 6
    styles['Heading3'].fontSize = 12
    styles['Heading3'].spaceAfter = 6
    styles['BodyText'].fontSize = 10
    styles['BodyText'].leading = 14
    styles['BodyText'].spaceAfter = 12

    # Add a custom style for references
    styles.add(ParagraphStyle(
        name='Reference',
        parent=styles['BodyText'],
        fontSize=10,
        leftIndent=36,
        firstLineIndent=-36,
        spaceAfter=12
    ))
    
    # Add a custom style for tables/monospace
    styles.add(ParagraphStyle(
        name='Monospace',
        parent=styles['BodyText'],
        fontName='Courier',
        fontSize=8,
        leading=10,
        spaceAfter=12
    ))

    story = []

    # Cover Page
    story.append(Paragraph("Deep Research AI Agent Report", styles['Title']))
    story.append(Spacer(1, 24))
    story.append(Paragraph(f"Query: {query}", styles['Normal']))
    story.append(Paragraph(f"Date: {datetime.date.today().strftime('%B %d, %Y')}", styles['Normal']))
    story.append(Paragraph(f"Author: [Your Name]", styles['Normal']))
    story.append(Spacer(1, 48))

    # Metadata
    story.append(Paragraph(f"OpenRouter Status: {'Operational' if check_openrouter_status() else 'Down'}", styles['Normal']))
    story.append(Spacer(1, 12))
    mode = "Deep Research" if deep_research else "Quick Research"
    story.append(Paragraph(f"Mode: {mode}", styles['Normal']))
    story.append(Spacer(1, 12))

    # Research Summary Section
    story.append(Paragraph("Research Summary", styles['Heading2']))
    story.append(Spacer(1, 12))

    # Add research data summary
    for item in data:
        story.append(Paragraph(f"‚Ä¢ {item['title']}", styles['Heading3']))
        story.append(Paragraph(item['content'], styles['BodyText']))
        story.append(Spacer(1, 12))

    # Process content based on format
    if parsed:
        # Handle JSON format
        for section in parsed.get("sections", []):
            title = section.get("title", "Section")
            content = section.get("content", "")
            
            # Add section title
            story.append(Paragraph(title, styles['Heading2']))
            story.append(Spacer(1, 12))
            
            # Clean and format content for PDF
            # Handle both ** and **** for bold (limit to reasonable length to avoid entire paragraphs)
            content = re.sub(r'\*{4}([^*]{1,100}?)\*{4}', r'<b>\1</b>', content)  # Convert **** to bold
            content = re.sub(r'\*\*([^*]{1,100}?)\*\*', r'<b>\1</b>', content)  # Convert ** to bold
            
            # Split into paragraphs
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    # Clean up any markdown artifacts
                    para = para.replace('\\n', ' ')
                    
                    lines = para.split('\n')
                    numbered_lines = [l for l in lines if re.match(r'^\d+\.\s', l.strip())]
                    
                    if "|" in para or "-|-" in para:
                        # Probably a table, use monospace
                        story.append(Paragraph(para.strip().replace("\n", "<br/>"), styles['Monospace']))
                        story.append(Spacer(1, 12))
                    elif len(numbered_lines) >= 2:  # It's a real list
                        # Split and add numbered items separately
                        for line in lines:
                            if line.strip():
                                story.append(Paragraph(line.strip(), styles['BodyText']))
                                if re.match(r'^\d+\.\s', line.strip()):
                                    story.append(Spacer(1, 12))  # Extra space after numbered items
                    else:
                        # Not a list, just a paragraph that might contain numbers
                        story.append(Paragraph(para.strip(), styles['BodyText']))
                        story.append(Spacer(1, 12))
        
        # Add references if present
        refs = parsed.get("references", [])
        if refs:
            story.append(PageBreak())
            story.append(Paragraph("References", styles['Heading2']))
            story.append(Spacer(1, 12))
            for i, ref in enumerate(refs, 1):
                ref_para = Paragraph(f"{i}. {ref}", styles['Reference'])
                story.append(ref_para)
    else:
        # Fallback to old text processing
        sections = summary.split("\n\n") if isinstance(summary, str) else []
        references = []
        current_heading = None
        
        for section in sections:
            if section.startswith("**") and section.endswith("**"):
                current_heading = section.strip("**").rstrip(":")
                if current_heading == "References":
                    continue
                story.append(Paragraph(current_heading, styles['Heading2']))
                story.append(Spacer(1, 12))
            else:
                if current_heading == "References":
                    refs = section.split("\n")
                    for ref in refs:
                        if ref.strip():
                            references.append(ref.strip())
                elif current_heading:
                    if "|" in section or "-|-" in section:
                        # Table in fallback path
                        story.append(Paragraph(section.strip().replace("\n", "<br/>"), styles['Monospace']))
                        story.append(Spacer(1, 12))
                    else:
                        formatted_text = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", section)
                        story.append(Paragraph(formatted_text, styles['BodyText']))
                        story.append(Spacer(1, 12))

        if references:
            story.append(PageBreak())
            story.append(Paragraph("References", styles['Heading2']))
            story.append(Spacer(1, 12))
            for i, ref in enumerate(references, 1):
                formatted_ref = format_reference_for_pdf(ref)
                if formatted_ref:
                    ref_para = Paragraph(f"{i}. {formatted_ref}", styles['Reference'])
                    story.append(ref_para)

    doc.build(story, onFirstPage=on_page, onLaterPages=on_page)
    buffer.seek(0)
    return buffer

# Helper function to add text with bold formatting to Word paragraphs
def _add_text_with_bold(paragraph, text):
    """Add text to a Word paragraph with bold markdown formatting."""
    import re
    # Split on bold markers, but limit to reasonable length spans
    parts = re.split(r'(\*\*[^*]{1,100}?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part:
            # Regular text
            paragraph.add_run(part)

# Function to generate Word document
def generate_docx(query, data, summary, deep_research=False):
    """Generate a Word document with query, data, and summary."""
    # Parse JSON if needed
    parsed = None
    if isinstance(summary, str) and summary.strip().startswith('{'):
        try:
            parsed = json.loads(summary)
        except:
            pass
    
    doc = Document()
    doc.add_heading("Deep Research AI Agent Report", 0)
    doc.add_paragraph(f"Query: {query}")
    doc.add_paragraph(f"Date: {datetime.date.today().strftime('%B %d, %Y')}")
    doc.add_paragraph(f"Author: Deep Research AI Agent")
    doc.add_paragraph(f"OpenRouter Status: {'Operational' if check_openrouter_status() else 'Down'}")
    doc.add_paragraph(f"Mode: {'Deep Research' if deep_research else 'Quick Research'}")
    
    # Research Summary Section
    doc.add_heading("Research Summary", level=1)
    if data:
        for item in data:
            p = doc.add_paragraph()
            p.add_run(f"‚Ä¢ {item['title']}").bold = True
            doc.add_paragraph(item['content'])
            doc.add_paragraph(f"Source: {item['url']}")
            doc.add_paragraph()  # Add spacing

    # Process content based on format
    if parsed:
        # Handle JSON format
        for section in parsed.get("sections", []):
            title = section.get("title", "Section")
            content = section.get("content", "")
            
            # Add section heading
            doc.add_heading(title, level=2)
            
            # First normalize **** to ** for consistent handling (limit to reasonable length)
            content = re.sub(r'\*{4}([^*]{1,100}?)\*{4}', r'**\1**', content)
            
            # Split into paragraphs
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    # Clean up any escaped newlines
                    para = para.replace('\\n', ' ')
                    # Check if this is a real numbered list (multiple numbered items)
                    lines = para.split('\n')
                    numbered_lines = [l for l in lines if re.match(r'^\d+\.\s', l.strip())]
                    
                    if len(numbered_lines) >= 2:  # It's a real list
                        # Add each line as a separate paragraph
                        for line in lines:
                            if line.strip():
                                p = doc.add_paragraph()
                                # Parse bold text
                                _add_text_with_bold(p, line.strip())
                                if re.match(r'^\d+\.\s', line.strip()):
                                    p.paragraph_format.space_after = Pt(18)  # Extra spacing for numbered items
                                else:
                                    p.paragraph_format.space_after = Pt(12)
                    else:
                        # Not a list, just a regular paragraph
                        p = doc.add_paragraph()
                        # Parse bold text
                        _add_text_with_bold(p, para.strip())
                        p.paragraph_format.space_after = Pt(12)
        
        # Add references
        refs = parsed.get("references", [])
        if refs:
            doc.add_page_break()
            doc.add_heading("References", level=2)
            for i, ref in enumerate(refs, 1):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.5)
                p.paragraph_format.space_after = Pt(12)
                p.add_run(f"{i}. {ref}")
    else:
        # Fallback to old text processing
        sections = summary.split("\n\n") if isinstance(summary, str) else []
        current_heading = None
        references = []
        
        for section in sections:
            if section.startswith("**") and section.endswith("**"):
                current_heading = section.strip("**").rstrip(":")
                if current_heading != "References":
                    doc.add_heading(current_heading, level=2)
            else:
                if current_heading == "References":
                    refs = section.split("\n")
                    for ref in refs:
                        if ref.strip():
                            references.append(ref.strip())
                elif current_heading:
                    formatted_text = re.sub(r"\*\*(.*?)\*\*", r"\1", section)
                    p = doc.add_paragraph(formatted_text)
                    p.paragraph_format.space_after = Pt(12)

        if references:
            doc.add_page_break()
            doc.add_heading("References", level=2)
            for i, ref in enumerate(references, 1):
                formatted_ref = format_reference_for_pdf(ref)
                if formatted_ref:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.first_line_indent = Inches(-0.5)
                    p.paragraph_format.space_after = Pt(12)
                    p.add_run(f"{i}. {formatted_ref}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Streamlit app setup
st.title("Deep Research AI Agent")
st.write("Enter a query to research and get a detailed response using Tavily and OpenRouter.")

# Load .env dynamically on each run to reflect external changes
try:
    load_dotenv(override=True)
except Exception:
    pass

# Sidebar: Status and Info
st.sidebar.header("OpenRouter Status")
if check_openrouter_status():
    st.sidebar.success("Operational", icon="‚úÖ")
else:
    st.sidebar.error("Down", icon="‚ùå")

st.sidebar.header("About")
st.sidebar.write("Dual-agent system using Tavily for research and OpenRouter for drafting with the model of your choosing.")
st.sidebar.write("Built with LangChain, LangGraph, and Streamlit.")

# ChromaDB Vector Store Toggle and Stats
st.sidebar.header("üß† Memory (ChromaDB)")

# Check if vector store is available
try:
    from vector_store import get_vector_store
    vector_store_available = True
except ImportError:
    vector_store_available = False
    st.sidebar.info("Install chromadb and sentence-transformers to enable memory", icon="‚ÑπÔ∏è")

if vector_store_available:
    # Toggle for enabling/disabling vector store
    enable_vector = st.sidebar.checkbox(
        "Enable Memory",
        value=os.getenv("ENABLE_VECTOR_STORE", "false").lower() == "true",
        help="Use local vector database to remember past research"
    )
    
    if enable_vector:
        # Set environment variable
        os.environ["ENABLE_VECTOR_STORE"] = "true"
        
        # Get vector store instance
        try:
            from vector_store import get_vector_store
            vs = get_vector_store()
            
            # Get stats for current language
            lang_map = {'English': 'en', 'Spanish': 'es', 'German': 'de'}
            current_lang = st.session_state.get('language', 'English')
            lang_code = lang_map.get(current_lang, 'en')
            
            stats = vs.get_stats(lang_code)
            
            # Display stats
            col1, col2 = st.sidebar.columns(2)
            with col1:
                st.metric("Active", stats['active'])
                st.metric("Sources", stats['sources'])
            with col2:
                st.metric("Expired", stats['expired'])
                st.metric("Summaries", stats['summaries'])
            
            # Show content type breakdown
            st.sidebar.caption(f"üì∞ News: {stats['news_content']} | üå≤ Evergreen: {stats['evergreen_content']}")
            
            # Management buttons
            col1, col2 = st.sidebar.columns(2)
            with col1:
                if st.button("Clear Expired", use_container_width=True):
                    cleared = vs.clear_expired()
                    st.sidebar.success(f"Cleared {cleared} expired items")
            with col2:
                if st.button("Clear All", use_container_width=True):
                    if vs.clear_all(lang_code):
                        st.sidebar.success("Memory cleared")
                    else:
                        st.sidebar.error("Failed to clear memory")
                        
        except Exception as e:
            st.sidebar.error(f"Vector store error: {str(e)[:100]}")
    else:
        os.environ["ENABLE_VECTOR_STORE"] = "false"
def _fix_readability(text: str) -> str:
    if not isinstance(text, str):
        return ""
    t = text.replace("\r\n", "\n").replace("\r", "\n")
    # Ensure a space after punctuation if followed by a letter (keep newlines)
    t = re.sub(r"([\.,;:])(?![\s\n])(?!$)(?=[A-Za-z])", r"\1 ", t)
    # Insert a space between ALL-CAPS token and following lowercase (avoid newlines)
    t = re.sub(r"\b([A-Z]{2,})([a-z])", r"\1 \2", t)
    # Insert a space between a closing alpha and opening parenthesis if missing
    t = re.sub(r"([A-Za-z])\(", r"\1 (", t)
    return t

def _render_small_text(text: str):
    t = _fix_readability(text)
    safe = html.escape(t)
    # Preserve paragraphs: split on blank lines
    parts = [p.strip() for p in safe.split("\n\n") if p.strip()]
    if not parts:
        st.markdown(f"<div style='font-size:0.95rem; line-height:1.7;'>{safe}</div>", unsafe_allow_html=True)
        return
    html_blocks = []
    for p in parts:
        p_clean = p.replace("\n", " ")
        html_blocks.append(f"<p style='margin:0 0 0.75rem 0;'>{p_clean}</p>")
    st.markdown("".join(html_blocks), unsafe_allow_html=True)

# Light cleanup for raw source content coming from the web
def _clean_source_text(text: str) -> str:
    if not isinstance(text, str):
        return ""
    # Normalize line endings
    t = text.replace("\r\n", "\n").replace("\r", "\n")
    
    # Drop common 'X min read' fragments
    t = re.sub(r"\b\d+\s*min\s*read\b", "", t, flags=re.IGNORECASE)
    
    # Replace lonely vertical bars used as separators with bullets
    t = re.sub(r"\s*\|\s*", " ‚Ä¢ ", t)
    
    # Ensure space after punctuation BUT NOT in decimal numbers (e.g. 4.6)
    # Match punctuation followed by non-space, only if not preceded AND followed by digits
    t = re.sub(r"(?<!\d)([\.,;:])(?!\s)|([\.,;:])(?!\s|(?<=\d.)\d)", r"\1\2 ", t)
    
    # Convert lone starting asterisks to proper bullet points for Markdown
    t = re.sub(r"^\s*\*\s*(?!\*)", "‚Ä¢ ", t, flags=re.MULTILINE)
    
    # Ensure no outer bolding wraps the entire block
    t = t.strip()
    if (t.startswith("**") and t.endswith("**")) or (t.startswith("__") and t.endswith("__")):
        if len(t) > 100:
            t = t[2:-2].strip()
    
    # Strip leading headers
    while t.startswith("#"):
        t = re.sub(r'^#+\s*', '', t).strip()
            
    return t


# API Keys gate & configuration
def _missing_api_keys():
    missing = []
    tav = os.getenv("TAVILY_API_KEY")
    router = os.getenv("OPENROUTER_API_KEY") or os.getenv("OPENAI_API_KEY")
    if not tav:
        missing.append("TAVILY_API_KEY")
    if not router:
        missing.append("OPENROUTER_API_KEY or OPENAI_API_KEY")
    return missing

@st.cache_data(ttl=300, show_spinner=False)
def _validate_openrouter_key(key: str) -> bool:
    try:
        if not key:
            return False
        headers = {"Authorization": f"Bearer {key}"}
        r = requests.get("https://openrouter.ai/api/v1/models", headers=headers, timeout=10)
        return r.status_code == 200
    except Exception:
        return False

@st.cache_data(ttl=300, show_spinner=False)
def _validate_tavily_key(key: str) -> bool:
    try:
        if not key:
            return False
        from tavily import TavilyClient
        client = TavilyClient(api_key=key)
        # Minimal request; if unauthorized will raise
        client.search("ping", max_results=1)
        return True
    except Exception:
        return False

def _persist_keys_to_env(env_updates: dict) -> bool:
    try:
        current = {}
        if os.path.exists(".env"):
            with open(".env", "r", encoding="utf-8") as f:
                for line in f:
                    if "=" in line and not line.strip().startswith("#"):
                        k, v = line.split("=", 1)
                        current[k.strip()] = v
        for k, v in env_updates.items():
            if v:
                current[k] = f"{k}={v}\n"
        with open(".env", "w", encoding="utf-8") as f:
            for _, line in current.items():
                f.write(line)
        return True
    except Exception as e:
        st.sidebar.error(f"Failed to write .env: {e}")
        return False

missing_keys = _missing_api_keys()

st.sidebar.header("API Keys")
if missing_keys:
    st.sidebar.warning("Missing: " + ", ".join(missing_keys))
    with st.sidebar.expander("Add API Keys", expanded=True):
        tavily = st.text_input("TAVILY_API_KEY", type="password", key="tavily_key_input")
        router = st.text_input("OPENROUTER_API_KEY or OPENAI_API_KEY", type="password", key="router_key_input")
        persist = st.checkbox("Save to .env", value=False, key="persist_keys_checkbox")
        if st.button("Save API Keys", key="save_keys_btn"):
            if tavily:
                os.environ["TAVILY_API_KEY"] = tavily
            if router:
                # Prefer OpenRouter key name; user can alternatively use OPENAI_API_KEY
                os.environ.setdefault("OPENAI_BASE_URL", os.getenv("OPENAI_BASE_URL") or "https://openrouter.ai/api/v1")
                os.environ["OPENROUTER_API_KEY"] = router
            if persist:
                updates = {}
                if tavily:
                    updates["TAVILY_API_KEY"] = tavily
                if router:
                    updates["OPENROUTER_API_KEY"] = router
                    updates.setdefault("OPENAI_BASE_URL", os.getenv("OPENAI_BASE_URL") or "https://openrouter.ai/api/v1")
                if _persist_keys_to_env(updates):
                    st.success("Saved to .env. Please rerun.")
            st.success("API keys saved for this session.")
            # Safe rerun across Streamlit versions
            try:
                st.rerun()
            except Exception:
                try:
                    st.experimental_rerun()
                except Exception:
                    pass
else:
    # Validate keys actually work against providers
    tav_ok = _validate_tavily_key(os.getenv("TAVILY_API_KEY",""))
    or_ok = _validate_openrouter_key(os.getenv("OPENROUTER_API_KEY") or os.getenv("OPENAI_API_KEY",""))
    if tav_ok and or_ok:
        st.sidebar.success("All required API keys detected.")
    else:
        st.sidebar.warning("API keys present but failed validation. Please re-enter.")
        if st.sidebar.button("Clear session API keys"):
            for k in ["TAVILY_API_KEY","OPENROUTER_API_KEY","OPENAI_API_KEY"]:
                if k in os.environ:
                    os.environ.pop(k)
            try:
                st.rerun()
            except Exception:
                pass

# -------- OpenRouter model selection (dynamic fetch with fallback) --------
@st.cache_data(ttl=600, show_spinner=False)
def fetch_openrouter_models(api_key: str | None) -> list:
    try:
        headers = {"Authorization": f"Bearer {api_key}"} if api_key else {}
        resp = requests.get("https://openrouter.ai/api/v1/models", headers=headers, timeout=15)
        if resp.status_code != 200:
            return []
        payload = resp.json()
        data = payload.get("data") or payload.get("models") or []
        models = []
        for m in data:
            mid = m.get("id") or m.get("name")
            if not mid:
                continue
            # Treat models marked ":free" or with zero pricing as free
            pricing = (m.get("pricing") or {})
            is_free = (":free" in mid) or all((not v or str(v).strip() in ("0", "0.0", "0.0000")) for v in pricing.values())
            # Attempt to capture context length from available fields
            ctx = m.get("context_length") or m.get("context")
            if not ctx:
                top = m.get("top_provider") or {}
                ctx = top.get("context_length") or top.get("context")
            models.append({"id": mid, "free": is_free, "context": ctx, "pricing": pricing})
        return models
    except Exception:
        return []

DEFAULT_MODEL = os.getenv("OPENROUTER_MODEL") or "tngtech/deepseek-r1t2-chimera:free"
api_key_present = os.getenv("OPENROUTER_API_KEY")
all_models = fetch_openrouter_models(api_key_present)
free_models_meta = [m for m in all_models if m.get("free")]
free_models = [m["id"] for m in free_models_meta]
if not free_models:
    # Fallback curated set
    free_models_meta = [
        {"id": "tngtech/deepseek-r1t2-chimera:free", "free": True, "context": None},
        {"id": "x-ai/grok-4-fast:free", "free": True, "context": None},
        {"id": "deepseek/deepseek-r1:free", "free": True, "context": None},
        {"id": "deepseek/deepseek-v3:free", "free": True, "context": None},
        {"id": "qwen/qwen2.5-7b-instruct:free", "free": True, "context": None},
    ]
    free_models = [m["id"] for m in free_models_meta]

st.sidebar.header("Model")
selected_model = st.sidebar.selectbox(
    "OpenRouter model (free preferred)",
    options=free_models,
    index=(free_models.index(DEFAULT_MODEL) if DEFAULT_MODEL in free_models else 0),
    help="Select a model to use for drafting."
)
os.environ["OPENROUTER_MODEL"] = selected_model

# Quick micro-benchmark for latency/throughput
@st.cache_data(ttl=900, show_spinner=False)
def benchmark_model(model_id: str, max_new_tokens: int = 64) -> dict | None:
    api_key = os.getenv("OPENROUTER_API_KEY")
    if not api_key:
        return None
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "http://localhost:8501",
        "X-Title": "Deep Research AI Agent"
    }
    body = {
        "model": model_id,
        "stream": True,
        "messages": [{"role": "user", "content": "Write a short sentence (~64 tokens)."}],
        "max_tokens": max_new_tokens,
        "temperature": 0.2
    }
    t0 = time.perf_counter()
    try:
        with requests.post(url, headers=headers, json=body, stream=True, timeout=60) as r:
            r.raise_for_status()
            ttfb = None
            token_chunks = 0
            for line in r.iter_lines(decode_unicode=True):
                if not line:
                    continue
                if not ttfb:
                    ttfb = time.perf_counter() - t0
                token_chunks += 1
        total = time.perf_counter() - t0
        throughput = token_chunks / max(total - (ttfb or 0), 1e-6)
        return {"latency_s": round(ttfb or total, 3), "throughput_tps": round(throughput, 1), "chunks": token_chunks}
    except Exception:
        return None

# Show model metrics UI
meta_map = {m["id"]: m for m in free_models_meta}
ctx_len = meta_map.get(selected_model, {}).get("context")
bench = None
if st.sidebar.button("Benchmark model", help="Measures time to first token and tokens/sec (short run)."):
    bench = benchmark_model(selected_model)
if bench:
    st.sidebar.metric("Latency", f"{bench['latency_s']} s")
    st.sidebar.metric("Throughput", f"{bench['throughput_tps']} tkn/s")
    st.sidebar.caption("Benchmarks are approximate and cached.")
elif ctx_len:
    st.sidebar.caption(f"Context: {ctx_len:,} tokens (reported)")
else:
    # If user clicked but benchmark failed
    if 'bench' in locals() and bench is None and api_key_present:
        st.sidebar.error("Benchmark failed (model/provider refused or network error). Try again or choose another model.")

# User input with Deep Research toggle
query = st.text_input("Research Query", "Latest advancements in quantum computing")
deep_research = st.checkbox("Deep Research Mode", value=False, help="Enable for a detailed, research-paper-style summary (5-6+ pages).")

# Research Settings Header
st.markdown("""
    <h2 style='color: #79c0ff; font-size: 1.8rem; font-weight: 600; margin-top: 1.8rem; margin-bottom: 1rem;'>
        üõ†Ô∏è Research Settings
    </h2>
""", unsafe_allow_html=True)

# Create three columns for better organization
col1, col2, col3 = st.columns(3)

with col1:
    st.markdown("#### Writing Style")
    writing_style = st.radio(
        "Select writing style",
        options=["Academic", "Business", "Technical", "Casual"],
        index=0,
        key="writing_style_radio",
        help="Choose the tone and style of your research"
    )

with col2:
    st.markdown("#### Language")
    language = st.radio(
        "Select language",
        options=["English", "Spanish", "German"],
        index=0,
        key="language_radio",
        help="Choose output language"
    )

with col3:
    st.markdown("#### Citation Format")
    citation_format = st.segmented_control(
        "Select citation style",
        options=["APA", "MLA", "IEEE", "BibTeX"],
        default="APA",
        key="citation_format_select",
        help="APA/MLA/IEEE are standard formats. BibTeX is for researchers using LaTeX."
    )

# Add word count slider below the columns
st.markdown("#### Target Word Count")
target_word_count = st.slider(
    "Select target word count",
    min_value=500,
    max_value=5000,
    value=1000,
    step=100,
    key="word_count_slider",
    help="Choose the approximate length of your research paper"
)

# Update all session state values
st.session_state.writing_style = writing_style
st.session_state.language = language
st.session_state.citation_format = citation_format
st.session_state.target_word_count = target_word_count

# Display current settings
st.markdown("""
    <div style='margin-top: 2rem;'>
        <p style='color: #58a6ff; font-size: 1.1rem; font-weight: 500;'>Current Settings:</p>
        <ul style='list-style-type: none; padding: 0; color: #c9d1d9;'>
            <li>üìù Writing Style: <strong>{}</strong></li>
            <li>üìö Citation Format: <strong>{}</strong></li>
            <li>üåê Language: <strong>{}</strong></li>
            <li>üìä Target Words: <strong>{}</strong></li>
        </ul>
    </div>
""".format(
    writing_style.title(),
    citation_format,
    language.title(),
    target_word_count
), unsafe_allow_html=True)

# Style descriptions dictionary with capitalized keys
style_descriptions = {
    "Academic": "Formal scholarly writing with rigorous citations",
    "Business": "Professional tone with actionable insights",
    "Technical": "Detailed technical analysis and specifications",
    "Casual": "Accessible, conversational explanation"
}

with st.expander("üìã Writing Style Preview"):
    st.markdown(f"""
        <div style='background-color: #161b22; border: 1px solid #30363d; border-radius: 8px; padding: 1rem;'>
            <p style='color: #58a6ff; font-weight: bold;'>{writing_style}</p>
            <p style='color: #c9d1d9;'>{style_descriptions[writing_style]}</p>
        </div>
    """, unsafe_allow_html=True)
    example_citation = {
        "title": "Impact of LLMs on Deep Research",
        "url": "https://arxiv.org/abs/2401.0001",
        "author": "Antigravity AI",
        "publisher": "Research DeepMind",
        "date": "2024"
    }
    st.code(format_citation(example_citation, citation_format))

# Cost Estimation Display
with st.container():
    est = estimate_research_cost(
        query=query, 
        deep_research=deep_research, 
        target_word_count=target_word_count,
        model=selected_model
    )
    
    # Color code confidence
    conf_colors = {"high": "#28a745", "medium": "#ffc107", "low": "#dc3545"}
    conf_color = conf_colors.get(est["confidence"], "#6c757d")
    
    st.markdown(f"""
        <div style='background-color: #161b22; border: 1px solid #30363d; border-radius: 8px; padding: 1rem; margin-bottom: 1rem;'>
            <div style='display: flex; justify-content: space-between; align-items: center;'>
                <span style='color: #8b949e; font-size: 0.9rem;'>Estimated Cost</span>
                <span style='color: {conf_color}; font-size: 0.8rem; font-weight: bold; border: 1px solid {conf_color}; padding: 2px 8px; border-radius: 12px;'>
                    {est['confidence'].upper()} CONFIDENCE
                </span>
            </div>
            <div style='font-size: 1.5rem; font-weight: bold; color: #e6edf3; margin: 0.5rem 0;'>
                ${est['cost_usd']['min']:.4f} - ${est['cost_usd']['max']:.4f}
            </div>
            <div style='color: #8b949e; font-size: 0.8rem;'>
                Tokens: {est['tokens']['min']:,} - {est['tokens']['max']:,} | Source: {est['pricing_source']}
            </div>
        </div>
    """, unsafe_allow_html=True)
    
    if est["degradation_reasons"]:
        with st.expander("Why is confidence lower?"):
            for reason in est["degradation_reasons"]:
                st.caption(f"‚Ä¢ {reason}")

# Research button logic (disabled if keys missing or busy)
run_disabled = (len(missing_keys) > 0) or st.session_state["__busy__"]
if len(missing_keys) > 0:
    st.info("Add your API keys to enable research.")

if st.button("Run Research", disabled=run_disabled):
    # Phase 1: mark busy and trigger the job, then rerun to render disabled state immediately
    st.session_state["__busy__"] = True
    st.session_state["__trigger__"] = True
    try:
        st.rerun()
    except Exception:
        pass

# Research execution loop (Granular Progress)
if st.session_state.get("__trigger__"):
    if not query.strip():
        st.error("Please enter a valid research query.")
        st.session_state["__busy__"] = False
        st.session_state["__trigger__"] = False
    else:
        try:
            from main import app as graph_app
            
            # Prepare inputs
            input_dict = {
                "query": query,
                "deep_research": deep_research,
                "target_word_count": target_word_count,
                "writing_style": writing_style.lower(),
                "citation_format": citation_format,
                "language": language.lower()
            }

            status_container = st.container()
            with status_container:
                status_box = st.status("üöÄ Initializing research workflow...", expanded=True)
            
            # Use app.stream for granular updates
            for event in graph_app.stream(input_dict):
                for node_name, output in event.items():
                    if node_name == "research":
                        num_found = len(output.get("research", []))
                        status_box.write(f"‚úÖ Phase 1/4: Found {num_found} sources.")
                        # Advance status label
                        status_box.update(label="Phase 2/4: Synthesizing & Drafting...", state="running")
                        time.sleep(0.01)
                    elif node_name == "draft":
                        status_box.write("‚úÖ Phase 2/4: Report text drafted.")
                        status_box.update(label="Phase 3/4: Finalizing documents...", state="running")
                        
                        # Store results from output
                        st.session_state.research_data = output.get("research", [])
                        st.session_state.response = output.get("draft", "")
                        
                        # Generate export formats
                        st.session_state.pdf_buffer = generate_pdf(query, st.session_state.research_data, st.session_state.response, deep_research=deep_research)
                        st.session_state.word_buffer = generate_docx(query, st.session_state.research_data, st.session_state.response, deep_research=deep_research)
                        
                        status_box.write("‚úÖ Phase 3/4: PDF & Word reports ready.")
                        status_box.update(label="Phase 4/4: Research Complete!", state="complete", expanded=False)
                        time.sleep(0.01)

            st.toast("Research successfully completed!", icon="üöÄ")
            st.success("Research completed! Explore your report below. üéâ")

        except Exception as e:
            st.error(f"Research failed: {str(e)}")
            logging.error(f"Research failed: {str(e)}")
        finally:
            st.session_state["__busy__"] = False
            st.session_state["__trigger__"] = False

# Display results
if st.session_state.research_data and st.session_state.response and not st.session_state.get("__busy__", False):
    # Re-display the summary that was generated
    try:
        parsed = json.loads(st.session_state.response)
    except Exception:
        parsed = {"sections": [{"title": "Response", "content": st.session_state.response}], "references": []}
    
    st.subheader("Structured Summary üìù")
    
    # Remove duplicate sections by tracking seen titles
    seen_titles = set()
    unique_sections = []
    for section in parsed.get("sections", []):
        title = section.get("title", "")
        if title not in seen_titles:
            seen_titles.add(title)
            unique_sections.append(section)
    
    st.write("### Research Data üìö")
    with st.expander("View External Source Snippets", expanded=True):
        for item in st.session_state.research_data:
            st.markdown(f"#### {item['title']}")
            # Clean and render content
            content = _clean_source_text(item['content'])
            
            # Use a slightly different style for raw snippets to distinguish from analysis
            st.info(content)
            
            st.markdown(f"üîó [Source]({item['url']})")
            st.markdown("---")
    
    st.write("### Detailed Analysis üìù")
    for section in unique_sections:
        title = section.get("title", "")
        content = section.get("content", "")
        
        with st.expander(f"üìñ {title}", expanded=True):
            # Deep clean the content for display
            display_content = _clean_analysis_text(content)
            
            # Normalize asterisk patterns for bolding
            display_content = re.sub(r'\*{4}([^*]{1,100}?)\*{4}', r'**\1**', display_content)
            display_content = re.sub(r'\*{3}([^*]{1,100}?)\*{3}', r'**\1**', display_content)
            
            # Use Streamlit markdown with unsafe HTML for tables if needed 
            # (Streamlit handles standard MD tables fine)
            st.markdown(display_content)

    # Show metadata
    meta = parsed.get("metadata", {})
    if meta.get("model"):
        st.caption(f"Model: {meta['model']} ‚Ä¢ Style: {meta.get('writing_style','-')} ‚Ä¢ Lang: {meta.get('language','-')} ‚Ä¢ Citations: {meta.get('citation_format','-')}")
    
    # Show references
    refs = parsed.get("references", [])
    if refs:
        with st.expander("References", expanded=True):
            for i, ref in enumerate(refs, 1):
                st.markdown(f"{i}. {ref}")
    
    # Show word count info
    word_count = len(st.session_state.response.split())
    page_estimate = word_count // 400 + 1
    st.info(f"Summary contains {word_count} words, estimated at {page_estimate} pages.")
    
    # Add Research Data section
    st.write("### Research Data üìö")
    for item in st.session_state.research_data:
        with st.expander(item['title']):
            # Render content with markdown support for all languages
            content = item['content']
            # Fix asterisk patterns
            content = re.sub(r'\*{4}([^*]{1,100}?)\*{4}', r'**\1**', content)
            content = re.sub(r'\*{3}([^*]{1,100}?)\*{3}', r'**\1**', content)
            st.markdown(content)
            st.markdown(f"[Visit Source]({item['url']})")

# Display download options if research data is available
if st.session_state.research_data and st.session_state.response:
    st.write("### Download Options")
    
    # Modern format selection
    selected_format = st.segmented_control(
        "Select format:",
        ["PDF (Recommended)", "Word", "Markdown", "JSON", "Text", "BibTeX"],
        default="PDF (Recommended)"
    )

    # Show description based on selection
    if selected_format == "PDF (Recommended)":
        st.caption("Download as a PDF file, ideal for printing or sharing with formatted content.")
        st.download_button(
            label="Download PDF üì•",
            data=st.session_state.pdf_buffer,
            file_name="research_report.pdf",
            mime="application/pdf"
        )
    elif selected_format == "BibTeX":
        st.caption("Download research citations in BibTeX format for LaTeX and reference managers.")
        # Generate BibTeX from citation_formatter
        try:
            from citation_formatter import CitationFormatter, Source
            sources = []
            for item in st.session_state.research_data:
                # Basic normalization for Source objects
                sources.append(Source(
                    title=item.get('title', 'Unknown Title'),
                    url=item.get('url', ''),
                    authors=[item.get('author')] if item.get('author') else None,
                    publisher=item.get('publisher'),
                    publication_date=None # We'll let from_dict handle this if needed better
                ))
            
            # CitationFormatter needs the list of sources in __init__
            formatter = CitationFormatter(sources)
            bib_data = formatter.format_bibtex()
        except Exception as e:
            bib_data = f"% Error generating BibTeX: {e}"
            
        st.download_button(
            label="Download BibTeX üì•",
            data=bib_data,
            file_name="citations.bib",
            mime="application/x-bibtex"
        )

    elif selected_format == "Word":
        st.caption("Download as a Word document, perfect for editing and professional documentation.")
        st.download_button(
            label="Download Word üì•",
            data=st.session_state.word_buffer,
            file_name="research_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    elif selected_format == "Markdown":
        st.caption("Download as a Markdown file, great for lightweight formatting and version control.")
        # Convert JSON response to pretty Markdown
        try:
            parsed = json.loads(st.session_state.response)
            lines = []
            lines.append(f"# {query}\n")
            for sec in parsed.get("sections", []):
                title = sec.get("title","Section")
                content = sec.get("content","")
                # Add extra newlines for numbered lists
                if re.search(r'\d+\.\s', content):
                    # Split and rejoin numbered items with double newlines
                    content = re.sub(r'(\d+\.\s)', r'\n\1', content)
                lines.append(f"\n## {title}\n{content}\n")
            refs = parsed.get("references", [])
            if refs:
                lines.append("\n## References\n")
                for i, r in enumerate(refs, 1):
                    lines.append(f"{i}. {r}")
            md_data = "\n".join(lines)
        except Exception:
            md_data = st.session_state.response
        st.download_button(
            label="Download Markdown üì•",
            data=md_data,
            file_name="research_summary.md",
            mime="text/markdown"
        )
    elif selected_format == "JSON":
        st.caption("Download as a JSON file (structured sections & references).")
        st.download_button(
            label="Download JSON üì•",
            data=st.session_state.response,
            file_name="research_summary.json",
            mime="application/json"
        )
    else:  # Text
        st.caption("Download as a plain text file, suitable for simple viewing or copying.")
        # Convert JSON to plain text
        try:
            parsed = json.loads(st.session_state.response)
            text_lines = []
            text_lines.append(f"RESEARCH SUMMARY\n{'='*50}\n")
            text_lines.append(f"Query: {query}\n")
            text_lines.append(f"Date: {datetime.date.today().strftime('%B %d, %Y')}\n")
            
            # Add metadata
            meta = parsed.get("metadata", {})
            if meta:
                text_lines.append(f"\nModel: {meta.get('model', 'N/A')}")
                text_lines.append(f"Writing Style: {meta.get('writing_style', 'N/A')}")
                text_lines.append(f"Language: {meta.get('language', 'N/A')}")
                text_lines.append(f"Citation Format: {meta.get('citation_format', 'N/A')}\n")
            
            text_lines.append(f"\n{'='*50}\n")
            
            # Add sections
            for section in parsed.get("sections", []):
                title = section.get("title", "Section")
                content = section.get("content", "")
                
                text_lines.append(f"\n{title.upper()}\n{'-'*len(title)}\n")
                
                # Clean content from markdown
                content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)
                
                # Format numbered lists properly
                if re.search(r'\d+\.\s', content):
                    # Add newlines before each number
                    content = re.sub(r'(\d+\.\s)', r'\n\1', content)
                
                text_lines.append(content)
                text_lines.append("\n")
            
            # Add references
            refs = parsed.get("references", [])
            if refs:
                text_lines.append(f"\nREFERENCES\n{'-'*10}\n")
                for i, ref in enumerate(refs, 1):
                    text_lines.append(f"{i}. {ref}\n")
            
            txt_data = "\n".join(text_lines)
        except Exception:
            # Fallback if JSON parsing fails
            txt_data = st.session_state.response
        
        st.download_button(
            label="Download Text üì•",
            data=txt_data,
            file_name="research_summary.txt",
            mime="text/plain"
        )

# Feedback form in sidebar
st.sidebar.header("Feedback")
feedback = st.sidebar.text_area("How can we improve?", placeholder="Share your thoughts...")
if st.sidebar.button("Submit Feedback"):
    if not feedback or len(feedback.strip()) == 0:
        st.sidebar.error("Please enter your feedback.")
    else:
        # Send feedback via server-side email
        try:
            import smtplib
            from email.mime.text import MIMEText
            from email.mime.multipart import MIMEMultipart
            
            # Get bot email credentials (sender and receiver are same)
            bot_email = os.getenv("FEEDBACK_BOT_EMAIL", "")
            bot_password = os.getenv("FEEDBACK_BOT_PASSWORD", "")
            
            if bot_email and bot_password:
                # Create message - bot sends to itself
                message = MIMEMultipart("alternative")
                message["Subject"] = f"Deep Research AI - User Feedback [{datetime.datetime.now().strftime('%Y-%m-%d')}]"
                message["From"] = bot_email
                message["To"] = bot_email
                
                # Create the email body
                text = f"""
                New feedback received from Deep Research AI Agent
                
                ========================================
                FEEDBACK:
                ========================================
                {feedback}
                
                ========================================
                Timestamp: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
                ========================================
                """
                
                part = MIMEText(text, "plain")
                message.attach(part)
                
                # Send email via Gmail SMTP (bot to itself)
                with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
                    server.login(bot_email, bot_password)
                    server.sendmail(bot_email, bot_email, message.as_string())
                
                st.sidebar.success("Thanks! Your feedback has been sent! üôè")
                
            else:
                # No feedback system configured
                st.sidebar.warning("Feedback system not available at the moment.")
                    
        except smtplib.SMTPAuthenticationError:
            st.sidebar.error("Feedback system authentication error. Please contact admin.")
        except Exception as e:
            st.sidebar.error(f"Unable to send feedback. Please try again later.")
